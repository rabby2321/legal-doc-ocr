from flask import Flask, request, jsonify
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
from concurrent.futures import ThreadPoolExecutor
from docx import Document
import os
import subprocess

app = Flask(__name__)

@app.route('/ocr', methods=['POST'])
def ocr():
    print("==== Incoming Request ====")
    print("FILES:", request.files)

    file = request.files.get('data')
    if not file:
        return jsonify({"error": "No file provided"}), 400

    filename = file.filename.lower()

    try:
        if filename.endswith(".pdf"):
            pdf_bytes = file.read()
            images = convert_from_bytes(pdf_bytes, dpi=120)

            with ThreadPoolExecutor() as executor:
                results = executor.map(pytesseract.image_to_string, images)

            text = "\n".join(results)
            return jsonify({"text": text.strip()})

        elif filename.endswith(".docx"):
            doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return jsonify({"text": text.strip()})

        elif filename.endswith(".doc"):
            temp_doc = "/tmp/upload.doc"
            temp_docx = "/tmp/upload.docx"
            file.save(temp_doc)

            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "docx", temp_doc, "--outdir", "/tmp"
            ], check=True)

            doc = Document(temp_docx)
            text = "\n".join([para.text for para in doc.paragraphs])

            os.remove(temp_doc)
            os.remove(temp_docx)

            return jsonify({"text": text.strip()})

        elif filename.endswith((".png", ".jpg", ".jpeg")):
            image = Image.open(file.stream)
            text = pytesseract.image_to_string(image)
            return jsonify({"text": text.strip()})

        else:
            return jsonify({"error": f"Unsupported file type: {filename}"}), 400

    except Exception as e:
        print("OCR ERROR:", str(e))
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
